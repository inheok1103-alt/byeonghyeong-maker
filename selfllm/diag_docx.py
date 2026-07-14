# -*- coding: utf-8 -*-
"""diag_docx.py — 진단 세트 JSON → 실물 시험지 스타일 DOCX(2단 꽉찬 레이아웃 + 약점 진단표 + 정답해설).
사용: python diag_docx.py <set.json> <out.docx>"""
import sys, io, json, re
from docx import Document
from docx.shared import Pt, Mm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

setFile = sys.argv[1] if len(sys.argv) > 1 else "selfllm/_diag_set.json"
outFile = sys.argv[2] if len(sys.argv) > 2 else "진단평가지.docx"
TITLE = sys.argv[3] if len(sys.argv) > 3 else "2026 여름특강 · 동화고 대비 영어 진단평가"
SET = json.load(io.open(setFile, encoding="utf-8"))
items = [q for q in SET.get("items", []) if q]
passage = SET.get("passage", "")
mcq = [q for q in items if q.get("choices") and len(q["choices"]) >= 4]
essay = [q for q in items if not q.get("choices") or len(q.get("choices", [])) < 4]
CIRC = ["", "①", "②", "③", "④", "⑤"]
FONT = "맑은 고딕"; SERIF = "Times New Roman"

def area(t):
    t = str(t or "")
    if re.search(r"어법|전환|도치|강조", t): return "어법·구문 정확성"
    if re.search(r"어휘|낱말|영영|함의|밑줄", t): return "어휘·문맥 판단"
    if re.search(r"빈칸|요약|주제|제목|요지|주장", t): return "논지·추론 독해"
    if re.search(r"내용|일치|지시|순서|삽입|무관|연결", t): return "정보·담화 파악"
    if re.search(r"서술|영작|조건|재배열|배열|우리말|수정|명시|찾아", t): return "서술형·영어 산출"
    return "종합"
# 약점 → 무엇을 못한 것 / 보강
WEAK = {
    "어휘·문맥 판단": ("글의 논지 방향으로 어휘를 판단하지 못함(뜻만 암기)", "담화방향 어휘·반의/혼동어 집중"),
    "어법·구문 정확성": ("수일치·시제·준동사·관계사 등 구문 정오 판단 약함", "문장 뼈대·동사형태 훈련"),
    "논지·추론 독해": ("본문 표현을 상위 개념으로 추상화·추론하지 못함", "요약어 추상화·빈칸 논지 훈련"),
    "정보·담화 파악": ("세부정보·지시어·담화 흐름을 정확히 못 잡음", "지시어 대상 고정·정보 대조 훈련"),
    "서술형·영어 산출": ("조건(단어수·제시어)에 맞춰 정확한 영어로 못 씀", "조건 재배열·형태변형·채점기준 훈련"),
}
AREAS = ["어휘·문맥 판단", "어법·구문 정확성", "논지·추론 독해", "정보·담화 파악", "서술형·영어 산출"]

def setfont(run, name=FONT, size=10, bold=False, color=None):
    run.font.name = name; run.font.size = Pt(size); run.font.bold = bold
    if color: run.font.color.rgb = color
    rPr = run._r.get_or_add_rPr(); rf = rPr.find(qn('w:rFonts'))
    if rf is None: rf = OxmlElement('w:rFonts'); rPr.append(rf)
    rf.set(qn('w:eastAsia'), name); rf.set(qn('w:ascii'), name); rf.set(qn('w:hAnsi'), name)

def para(container, after=2, before=0, space=1.0, align=None):
    p = container.add_paragraph()
    pf = p.paragraph_format; pf.space_after = Pt(after); pf.space_before = Pt(before)
    pf.line_spacing = space
    if align: p.alignment = align
    return p

def run(p, text, size=10, bold=False, name=FONT, color=None, ul=False):
    r = p.add_run(text); setfont(r, name, size, bold, color); r.font.underline = ul; return r

def run_html(p, text, size=10, name=SERIF, color=None):
    """<u>밑줄</u>·<b>굵게</b> 태그를 실제 서식으로, 그 외 태그는 제거."""
    text = str(text or "")
    ul = False; bold = False
    for tok in re.split(r'(</?[ub]>)', text):
        if tok == '<u>': ul = True; continue
        if tok == '</u>': ul = False; continue
        if tok == '<b>': bold = True; continue
        if tok == '</b>': bold = False; continue
        if not tok: continue
        clean = re.sub(r'<[^>]+>', '', tok)
        if clean: run(p, clean, size, bold=bold, name=name, color=color, ul=ul)

def set_cols(section, num=2, sep=True, space_mm=6):
    cols = section._sectPr.find(qn('w:cols'))
    if cols is None: cols = OxmlElement('w:cols'); section._sectPr.append(cols)
    cols.set(qn('w:num'), str(num)); cols.set(qn('w:sep'), '1' if sep else '0')
    cols.set(qn('w:space'), str(int(space_mm * 56.7)))

def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr(); sh = OxmlElement('w:shd')
    sh.set(qn('w:val'), 'clear'); sh.set(qn('w:fill'), hexcolor); tcPr.append(sh)

def no_space_after(p): p.paragraph_format.space_after = Pt(0)

def keep(p, with_next=True):
    """문항이 컬럼·페이지 경계에서 잘리지 않게: 줄 뭉침 + 다음 단락과 붙임."""
    pPr = p._p.get_or_add_pPr()
    for tag in (["w:keepLines", "w:keepNext"] if with_next else ["w:keepLines"]):
        if pPr.find(qn(tag)) is None:
            el = OxmlElement(tag); el.set(qn('w:val'), '1'); pPr.append(el)
    return p

doc = Document()
# 여백을 좁혀 꽉 차게
sec = doc.sections[0]
sec.top_margin = Mm(11); sec.bottom_margin = Mm(11); sec.left_margin = Mm(12); sec.right_margin = Mm(12)
sec.page_height = Mm(297); sec.page_width = Mm(210)

# ── 시험지 헤더 표(동화고 실물 스타일) : 1단 구간 ──
title = para(doc, after=3, align=WD_ALIGN_PARAGRAPH.CENTER)
run(title, TITLE, 15, True)
sub = para(doc, after=5, align=WD_ALIGN_PARAGRAPH.CENTER)
run(sub, "고1 공통영어 · 제한시간 40분 · 선택형 %d + 논술형 %d = %d문항" % (len(mcq), len(essay), len(items)), 9, False, color=RGBColor(0x55, 0x55, 0x55))

# 인적사항 표
t = doc.add_table(rows=1, cols=6); t.style = "Table Grid"
hdrs = [("학교 / 반", ""), ("이름", ""), ("점수", "        / 100")]
cells = t.rows[0].cells
labels = ["학교·반", "이름", "점수"]
for i in range(3):
    c1 = cells[i*2]; c2 = cells[i*2+1]
    c1.width = Mm(22); c2.width = Mm(46)
    shade(c1, "F0EFE9")
    p1 = c1.paragraphs[0]; p1.alignment = WD_ALIGN_PARAGRAPH.CENTER; run(p1, labels[i], 9, True); no_space_after(p1)
    p2 = c2.paragraphs[0]; run(p2, "        / 100" if i == 2 else "", 9); no_space_after(p2)
for row in t.rows:
    row.height = Mm(8)

g = para(doc, before=4, after=2)
run(g, "◈ 안내  ", 8.5, True, color=RGBColor(0xB4, 0x84, 0x2A))
run(g, "이 지문은 교과서에 없는 '처음 보는 외부지문'입니다. 동화고 시험처럼, 낯선 지문을 유형별로 대응하는 힘을 봅니다. 모르는 문항은 넘기지 말고 여백에 '왜 어려운지'를 표시하세요. 서술형은 조건(단어 수·제시어)을 지켜야 부분점수를 받습니다.", 8.5, color=RGBColor(0x4a, 0x3c, 0x1a))

# ── 2단 구간 시작(문제) ──
doc.add_section(WD_SECTION.CONTINUOUS)
sec2 = doc.sections[-1]
sec2.top_margin = Mm(11); sec2.bottom_margin = Mm(11); sec2.left_margin = Mm(12); sec2.right_margin = Mm(12)
set_cols(sec2, 2, True, 6)

no = 0
def add_passage(container, text):
    p = para(container, after=3, before=1, space=1.15)
    run_html(p, text, 9, name=SERIF)   # <u> 밑줄 실제 서식 적용
    p.paragraph_format.left_indent = Mm(1); p.paragraph_format.right_indent = Mm(1)
    return p

def strip_markup(s):
    s = re.sub(r'</?[ub]>', '', str(s or ''))
    s = re.sub(r'_{2,}', '', s)
    return re.sub(r'\s+', ' ', s).strip()
CLEAN = strip_markup(passage)

# 공통 지문을 상단에 한 번만 제시(단일지문 진단). 다지문 동형모고(SET.passage="")면 문항별 지문 인쇄.
if CLEAN:
    lead = para(doc, after=1, before=2)
    run(lead, "[%d~%d] 다음 글을 읽고 물음에 답하시오. (어법·어휘·빈칸 문항은 각 문항의 밑줄·빈칸 지문을 따름)" % (1, len(items)), 9, True, color=RGBColor(0x16, 0x23, 0x3F))
    lp = para(doc, after=4, before=1, space=1.15)
    run(lp, CLEAN, 9, name=SERIF); lp.paragraph_format.left_indent = Mm(1); lp.paragraph_format.right_indent = Mm(1)

for q in mcq:
    no += 1
    blk = []   # 이 문항의 모든 단락 — keep으로 함께 묶어 잘림 방지
    p = para(doc, after=1, before=3, space=1.05)
    run(p, "%d. " % no, 10, True, color=RGBColor(0xB4, 0x84, 0x2A))
    run(p, str(q.get("instruction", "")), 9.5, True); blk.append(p)
    pg = q.get("passage") or passage
    # 표식형(밑줄·빈칸으로 원문과 달라진 경우)만 지문 표시. 공통 clean 지문과 같으면 상단 참조.
    if pg and strip_markup(pg) != CLEAN:
        blk.append(add_passage(doc, pg))
    elif pg:
        rp = para(doc, after=1, space=1.05); rp.paragraph_format.left_indent = Mm(1)
        run(rp, "(윗글 참조)", 8.5, color=RGBColor(0x88, 0x88, 0x88)); blk.append(rp)
    for i, c in enumerate(q["choices"]):
        cp = para(doc, after=0, space=1.05); cp.paragraph_format.left_indent = Mm(1)
        run(cp, CIRC[i+1] + " ", 9.3, name=SERIF); run_html(cp, str(c), 9.3, name=SERIF); blk.append(cp)
    for j, bp in enumerate(blk): keep(bp, with_next=(j < len(blk) - 1))

if essay:
    sp = para(doc, before=4, after=3)
    r = run(sp, "  서술형  ", 10, True); shade_run = None
    # 서술형 구획 표시(테두리 대신 배경 런)
    rpr = sp.runs[0]._r.get_or_add_rPr(); sh = OxmlElement('w:shd'); sh.set(qn('w:val'), 'clear'); sh.set(qn('w:fill'), '1A2030'); rpr.append(sh)
    sp.runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for q in essay:
        no += 1
        blk = []
        ins = str(q.get("instruction", "")); cond = ""
        ci = ins.find("〈조건〉");
        if ci < 0: ci = ins.find("[조건]")
        if ci >= 0: cond = ins[ci:]; ins = ins[:ci]
        p = para(doc, after=1, before=3, space=1.05)
        run(p, "%d. " % no, 10, True, color=RGBColor(0xB4, 0x84, 0x2A)); run(p, ins.strip(), 9.5, True); blk.append(p)
        pgE = q.get("passage") or passage
        if pgE and strip_markup(pgE) != CLEAN:
            blk.append(add_passage(doc, pgE))
        if cond:
            cpar = para(doc, after=2, space=1.1); cpar.paragraph_format.left_indent = Mm(1)
            run(cpar, cond.strip(), 8.7, color=RGBColor(0x33, 0x33, 0x33)); blk.append(cpar)
        # 답란
        ap = para(doc, after=4, space=1.4); run(ap, "답: ", 9, True)
        ap.paragraph_format.space_after = Pt(18); blk.append(ap)
        for j, bp in enumerate(blk): keep(bp, with_next=(j < len(blk) - 1))

# ── 정답·해설·약점 진단(새 페이지, 1단) ──
doc.add_section(WD_SECTION.NEW_PAGE)
sec3 = doc.sections[-1]
sec3.top_margin = Mm(12); sec3.bottom_margin = Mm(11); sec3.left_margin = Mm(13); sec3.right_margin = Mm(13)
set_cols(sec3, 1)

h = para(doc, after=6, align=WD_ALIGN_PARAGRAPH.CENTER); run(h, "정답 · 약점 진단 · 해설", 14, True)

# 정답표
run(para(doc, after=2, before=2), "■ 객관식 정답", 11, True, color=RGBColor(0xB4, 0x84, 0x2A))
at = doc.add_table(rows=2, cols=len(mcq)); at.style = "Table Grid"
for i, q in enumerate(mcq):
    c0 = at.rows[0].cells[i]; c1 = at.rows[1].cells[i]
    shade(c0, "F0EFE9")
    p0 = c0.paragraphs[0]; p0.alignment = WD_ALIGN_PARAGRAPH.CENTER; run(p0, str(i+1), 9, True); no_space_after(p0)
    p1 = c1.paragraphs[0]; p1.alignment = WD_ALIGN_PARAGRAPH.CENTER; run(p1, CIRC[q.get("answer", 0)] if q.get("answer") else "-", 11, True); no_space_after(p1)

# 약점 진단표: 틀린 문항 → 약점 → 무엇을 못한 것 → 보강
run(para(doc, after=2, before=8), "■ 무엇을 틀리면 무엇이 약점인가 (틀린 문항 번호로 약점 확인)", 11, True, color=RGBColor(0xB4, 0x84, 0x2A))
# 영역별 문항 매핑
byArea = {a: [] for a in AREAS}
n = 0
for q in items:
    n += 1; a = area(q.get("type")); byArea.setdefault(a, []).append(n)
wt = doc.add_table(rows=1, cols=5); wt.style = "Table Grid"
whead = ["약점 영역", "이 문항 틀리면", "무엇을 못한 것", "여름특강 보강", "내 체크"]
wmm = [30, 22, 62, 46, 18]
for i, htext in enumerate(whead):
    c = wt.rows[0].cells[i]; c.width = Mm(wmm[i]); shade(c, "16233F")
    pp = c.paragraphs[0]; pp.alignment = WD_ALIGN_PARAGRAPH.CENTER; run(pp, htext, 8.6, True, color=RGBColor(0xFF, 0xFF, 0xFF)); no_space_after(pp)
for a in AREAS:
    nums = byArea.get(a, [])
    if not nums: continue
    row = wt.add_row().cells
    row[0].width = Mm(wmm[0]); shade(row[0], "FAF8F2")
    p = row[0].paragraphs[0]; run(p, a, 8.8, True); no_space_after(p)
    p = row[1].paragraphs[0]; run(p, ", ".join(str(x) for x in nums) + "번", 8.8, True, color=RGBColor(0xC0, 0x26, 0x5A)); no_space_after(p)
    miss, fix = WEAK.get(a, ("", ""))
    p = row[2].paragraphs[0]; run(p, miss, 8.6); no_space_after(p)
    p = row[3].paragraphs[0]; run(p, fix, 8.6, color=RGBColor(0x0d, 0x5c, 0x3f)); no_space_after(p)
    p = row[4].paragraphs[0]; run(p, "", 8.6); no_space_after(p)

dp = para(doc, before=5, after=4)
run(dp, "◈ ", 9, True, color=RGBColor(0xB4, 0x84, 0x2A))
run(dp, "틀린 문항이 몰린 영역이 이번 여름특강의 1순위 보강 지점입니다. 6회 커리(문장뼈대→동사형태→절연결→어휘·지시어→서술형 재배열→실전 리포트)에서 그 영역부터 배정합니다.", 9, color=RGBColor(0x4a, 0x3c, 0x1a))
# 평가 방향(실물 기말 근거)
vp = para(doc, before=2, after=3)
run(vp, "◈ 평가 방향(동화고 실물 기말 근거)  ", 8.8, True, color=RGBColor(0x16, 0x23, 0x3F))
run(vp, "동화고 기말은 선택형16+논술형5(논술형 39%)로, 같은 지문을 어법·어휘·빈칸·지시어·요약으로 바꿔 묻고 논술형은 '조건에 맞춘 정확한 영어 산출'을 채점합니다. 특히 ①문맥상 어휘를 '논지 방향'으로 판단, ②수일치·준동사·관계사 복합 어법, ③본문 표현을 상위개념으로 추상화, ④조건부 영작(단어수·제시어·형태변화)이 승부처입니다. 이 진단지는 그 4개 축을 그대로 반영해 약점을 잡습니다.", 8.8, color=RGBColor(0x33, 0x33, 0x33))
# 예고 없는 외부지문 대비(동화고 최대 관건)
xp = para(doc, before=2, after=6)
run(xp, "◈ 핵심 관건 — 예고 없는 외부지문  ", 8.8, True, color=RGBColor(0xC0, 0x26, 0x5A))
run(xp, "동화고의 진짜 난도는 '시험 당일 처음 보는 외부지문'에 있습니다(교과서 밖 학술·저널 지문을 예고 없이 출제). 그래서 대비는 특정 지문 암기가 아니라 '어떤 낯선 지문이 와도 유형별로 대응하는 전이 가능한 능력'을 키우는 것입니다. 이 진단지도 처음 보는 외부지문으로 그 대응력을 측정합니다. Ray 시스템은 무한한 '새 외부지문'으로 훈련시켜, 시험의 예고 없는 외부지문이 학생에게는 '이미 수십 개 풀어 본 유형 중 하나'가 되게 합니다. → 매 수업 낯선 외부지문 1편으로 4축(어휘 방향·복합 어법·논지 추상화·조건 영작)을 실전 훈련.", 8.8, color=RGBColor(0x33, 0x33, 0x33))

# 문항별 해설
run(para(doc, after=3, before=6), "■ 문항별 정답·해설", 11, True, color=RGBColor(0xB4, 0x84, 0x2A))
n = 0
for q in items:
    n += 1
    exp = str(q.get("explanation", ""))
    model = ""
    if "[모범답안]" in exp:
        model = exp.split("[모범답안]")[1].split("[채점")[0].split("【")[0].strip()
    ansline = ""
    if q.get("choices") and len(q["choices"]) >= 4 and q.get("answer"):
        ansline = CIRC[q["answer"]] + " " + str(q["choices"][q["answer"]-1])
    else:
        ansline = model[:80] if model else "(서술형)"
    p = para(doc, after=1, before=4, space=1.05)
    run(p, "%d. [%s · %s] " % (n, q.get("type"), area(q.get("type"))), 9.5, True)
    run(p, "정답 " + ansline, 9.5, True, color=RGBColor(0x0d, 0x5c, 0x3f))
    # 해설 본문(오답진단 앞부분까지)
    body = exp.split("[오답 진단")[0].strip()
    body = re.sub(r"【출제의도】.*", "", body).strip()
    if body:
        bp = para(doc, after=2, space=1.1); bp.paragraph_format.left_indent = Mm(2)
        run(bp, body[:400], 8.8, color=RGBColor(0x33, 0x33, 0x33))

doc.save(outFile)
print("SAVED", outFile)
